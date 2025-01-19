from docx import Document
from .base_step import Step

class DocumentReaderStep(Step):
    def __init__(self):
        self.texts = []
        pass

    async def process(self, file_path: str):
        """
        读取Word文档内容
        """
        try:
            if not isinstance(file_path, str):
                raise ValueError(f"Expected string file path, got {type(file_path)}")
                
            doc = Document(file_path)
            full_text = []
            
            # 读取所有段落
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:  # 忽略空段落
                    full_text.append(text)
                    
            # 读取所有表格中的文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text:  # 忽略空单元格
                            full_text.append(text)
            
            # 合并所有文本
            self.texts.append("\n".join(full_text))
            return self.texts
            
        except Exception as e:
            raise Exception(f"Error reading document: {str(e)}")